"""DOCX Template Renderer and PDF Converter"""
import os
from pathlib import Path
from datetime import datetime
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import logging

logger = logging.getLogger(__name__)

BACKEND_DIR = Path(__file__).parent.parent
TEMPLATES_DIR = BACKEND_DIR / 'templates'
TEMPS_DIR = BACKEND_DIR / 'temps'

# Create temps directory if not exists
TEMPS_DIR.mkdir(exist_ok=True)


from datetime import datetime
import re

ID_MONTHS = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April",
    5: "Mei", 6: "Juni", 7: "Juli", 8: "Agustus",
    9: "September", 10: "Oktober", 11: "November", 12: "Desember"
}

def format_indonesian_date(date_str):
    bulan = {
        "01": "Januari", "02": "Februari", "03": "Maret", "04": "April",
        "05": "Mei", "06": "Juni", "07": "Juli", "08": "Agustus",
        "09": "September", "10": "Oktober", "11": "November", "12": "Desember",
    }

    # Input pasti format YYYY-MM-DD dari HTML <input type="date">
    try:
        tahun = date_str[:4]
        bulan_num = date_str[5:7]
        hari = str(int(date_str[8:10]))  # Buang leading zero
        return f"{hari} {bulan[bulan_num]} {tahun}"
    except:
        return date_str


def format_id_date(date_input):
    """
    Convert various date inputs to Indonesian long date format: e.g. "2 Desember 2025".
    Accepts:
      - datetime.datetime or datetime.date
      - strings like "02/12/2025", "2/12/2025", "2025-12-02", "2025/12/02"
    If parsing fails, returns the original input (string).
    """
    if not date_input:
        return ""

    # if it's already a datetime/date
    if isinstance(date_input, datetime):
        dt = date_input
    else:
        # if it's bytes or other, convert to str
        s = str(date_input).strip()

        # Try common formats
        tried = [
            "%d/%m/%Y",
            "%d-%m-%Y",
            "%Y-%m-%d",
            "%Y/%m/%d",
            "%d %m %Y",
            "%d %B %Y",  # e.g. "2 December 2025"
            "%d %b %Y",
        ]
        dt = None
        for fmt in tried:
            try:
                dt = datetime.strptime(s, fmt)
                break
            except Exception:
                continue

        # If still not parsed, try to extract numbers with regex (like 02/12/2025)
        if dt is None:
            m = re.search(r'(\d{1,4})\D+(\d{1,2})\D+(\d{1,4})', s)
            if m:
                a, b, c = m.group(1), m.group(2), m.group(3)
                # heuristik: if first part has 4 digits => YYYY MM DD pattern
                try:
                    if len(a) == 4:
                        y = int(a); mo = int(b); day = int(c)
                    else:
                        day = int(a); mo = int(b); y = int(c)
                    dt = datetime(year=y, month=mo, day=day)
                except Exception:
                    dt = None

        if dt is None:
            # parsing gagal — kembalikan original string
            return s

    # format as "2 Desember 2025" (no leading zero for day)
    day = dt.day
    month_name = ID_MONTHS.get(dt.month, dt.strftime("%B"))
    year = dt.year
    return f"{day} {month_name} {year}"



def render_docx_with_docxtpl(data: dict, logo_path: str = None) -> str:
    """
    Render a DOCX template using docxtpl with patient data and logo
    """
    try:
        from docxtpl import InlineImage
        from docx.shared import Mm

        # Ambil pilihan ukuran kertas dari user
        paper_size = data.get("paper_size", "A4")

        # Tentukan template berdasarkan pilihan user
        if paper_size == "A3":
            template_path = TEMPLATES_DIR / "sick_letter_horizontal.docx"
        elif paper_size == "A4":
            template_path = TEMPLATES_DIR / "sick_letter.docx"
        elif paper_size == "A3Sehat":
            template_path = TEMPLATES_DIR / "sick_letter_horizontal_sehat.docx"
        elif paper_size == "A4Sehat":
            template_path = TEMPLATES_DIR / "sick_letter_sehat.docx"
        else:
            # fallback kalau nilainya aneh / tidak dikenal
            template_path = TEMPLATES_DIR / "sick_letter.docx"


        # Debug log
        logger.info(f"Using template: {template_path}")

        # ambil raw value (bisa string "02/12/2025" atau None)
        raw_date_issued = data.get('date_issued')
        if not raw_date_issued:
            raw_date_issued = datetime.now()

        formatted_date_issued = format_id_date(raw_date_issued)
        
        context = {
            'name': data.get('name', ''),
            'gender': data.get('gender', ''),
            'age': data.get('age', ''),
            'occupation': data.get('occupation', ''),
            'address': data.get('address', ''),
            'height': data.get('height', ''),
            'weight': data.get('weight', ''),
            'duration': data.get('duration', ''),
            'from_date': data.get('from_date', ''),
            'to_date': data.get('to_date', ''),
            'notes': data.get('notes', ''),
            'clinic_name': data.get('clinic_name', ''),
            'clinic_type': data.get('clinic_type', ''),
            'clinic_address': data.get('clinic_address', ''),
            'doctor_name': data.get('doctor_name', 'dr. [Nama Dokter]'),
            'doctor_nip': data.get('doctor_nip', '[NIP Dokter]'),
            'letter_number': data.get('letter_number', f"SKS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"),
            'date_issued': formatted_date_issued,
            'location': data.get('location', 'Jakarta')
        }
        
        doc = DocxTemplate(template_path)

        # Logo
        if logo_path and os.path.exists(logo_path):
            context['logo'] = InlineImage(doc, logo_path, width=Mm(56.5), height=Mm(51.5))
        else:
            context['logo'] = ""

        doc.render(context)

        output_filename = f"sick_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = TEMPS_DIR / output_filename
        doc.save(output_path)

        logger.info(f"DOCX rendered with template {template_path}: {output_path}")
        return str(output_path)

    except Exception as e:
        logger.error(f"Error rendering DOCX: {str(e)}")
        raise


def create_docx_from_scratch(data: dict, logo_path: str = None, paper_size: str = 'A4') -> str:
    """
    Create a DOCX document from scratch (fallback if template doesn't exist)
    
    Args:
        data: Dictionary containing patient and clinic information
        logo_path: Path to the clinic logo file
        paper_size: Paper size (A4 or A5)
    
    Returns:
        Path to the created DOCX file
    """
    try:
        doc = Document()
        
        # Set paper size
        section = doc.sections[0]
        if paper_size.upper() == 'A5':
            section.page_height = Inches(8.27)  # 210mm
            section.page_width = Inches(5.83)   # 148mm
        else:  # A4
            section.page_height = Inches(11.69)  # 297mm
            section.page_width = Inches(8.27)    # 210mm
        
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header with logo and clinic info
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        
        # Left cell - Clinic info
        left_cell = header_table.rows[0].cells[0]
        left_paragraph = left_cell.paragraphs[0]
        
        clinic_name = data.get('clinic_name', 'Dinas Kesehatan')
        clinic_type = data.get('clinic_type', 'Puskesmas')
        clinic_address = data.get('clinic_address', '')
        
        run = left_paragraph.add_run(f"{clinic_name}\n")
        run.font.size = Pt(12)
        run.font.bold = True
        
        run = left_paragraph.add_run(f"{clinic_type}\n")
        run.font.size = Pt(11)
        
        run = left_paragraph.add_run(clinic_address)
        run.font.size = Pt(9)
        
        # Right cell - Logo
        if logo_path and os.path.exists(logo_path):
            right_cell = header_table.rows[0].cells[1]
            right_paragraph = right_cell.paragraphs[0]
            right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            right_paragraph.add_run().add_picture(logo_path, width=Inches(1.2))
        
        # Separator line
        doc.add_paragraph('_' * 80)
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('SURAT KETERANGAN DOKTER')
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        # Letter number
        letter_num = doc.add_paragraph()
        letter_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        letter_number = data.get('letter_number', f'SKS-{datetime.now().strftime("%Y%m%d-%H%M%S")}')
        num_run = letter_num.add_run(f'No: {letter_number}')
        num_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # Opening statement
        doc.add_paragraph('Yang bertanda tangan di bawah ini menerangkan bahwa:')
        
        # Patient biodata
        biodata_items = [
            ('Nama', data.get('name', '')),
            ('Jenis Kelamin', data.get('gender', '')),
            ('Umur', f"{data.get('age', '')} tahun"),
            ('Pekerjaan', data.get('occupation', '')),
            ('Alamat', data.get('address', '')),
            ('Tinggi Badan', data.get('height', '')),
            ('Berat Badan', data.get('weight', '')),
        ]
        
        for label, value in biodata_items:
            p = doc.add_paragraph()
            p.add_run(f"{label:<20}").font.size = Pt(11)
            p.add_run(': ').font.size = Pt(11)
            run = p.add_run(value)
            run.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
        
        # Medical statement
        statement = doc.add_paragraph(
            f"Bahwa pada pemeriksaan kesehatan saat ini, pasien tersebut dalam keadaan sakit dan "
            f"membutuhkan istirahat selama {data.get('duration', '')} hari, "
            f"mulai tanggal {data.get('from_date', '')} sampai {data.get('to_date', '')}."
        )
        statement.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Diagnosis notes
        if data.get('notes'):
            doc.add_paragraph()  # Spacing
            notes_p = doc.add_paragraph()
            notes_p.add_run('Catatan Diagnosa: ').font.bold = True
            notes_p.add_run(data.get('notes'))
        
        doc.add_paragraph()  # Spacing
        
        # Closing statement
        doc.add_paragraph(
            'Demikian surat keterangan ini dibuat dengan sebenarnya untuk dapat dipergunakan sebagaimana mestinya.'
        )
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph()  # Spacing
        
        # Signature block
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.autofit = False
        
        # Right cell for signature
        right_cell = sig_table.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        location = data.get('location', 'Jakarta')
        raw_date_issued = data.get('date_issued')
        if not raw_date_issued:
            raw_date_issued = datetime.now()
        date_issued = format_id_date(raw_date_issued)
        right_para.add_run(f"{location}, {date_issued}\n\n\n\n")
        
        doctor_name = data.get('doctor_name', 'dr. [Nama Dokter]')
        doctor_nip = data.get('doctor_nip', '[NIP Dokter]')
        
        doc_run = right_para.add_run(f"{doctor_name}\n")
        doc_run.font.bold = True
        right_para.add_run(f"NIP. {doctor_nip}")
        
        # Save document
        output_filename = f"sick_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = TEMPS_DIR / output_filename
        doc.save(str(output_path))
        
        logger.info(f"DOCX created from scratch successfully: {output_path}")
        return str(output_path)
        
    except Exception as e:
        logger.error(f"Error creating DOCX from scratch: {str(e)}")
        raise
